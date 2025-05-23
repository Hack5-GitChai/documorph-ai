# app/services/doc_formatter_v1.py

from pathlib import Path
from docx import Document 
# Removed HTTPException as services shouldn't typically raise HTTP specific exceptions directly
# They should raise custom exceptions or standard Python exceptions that routes can handle.

async def apply_dummy_formatting(
    input_file_path_str: str, 
    output_file_path_str: str
    # template_path_str: str, # We'll add this back when we use real templates
) -> str:
    """
    Simulates formatting a document (reads from input, writes a dummy DOCX to output).
    For this dummy version, it assumes the input is a text file.

    Args:
        input_file_path_str (str): The absolute or relative path to the input file.
        output_file_path_str (str): The absolute or relative path where the output .docx should be saved.

    Returns:
        str: The path to the created output .docx file.

    Raises:
        FileNotFoundError: If the input file does not exist (if critical).
        IOError: If there's an issue reading the input or writing the output file.
        Exception: For other general errors during DOCX creation.
    """
    input_path = Path(input_file_path_str)
    output_path = Path(output_file_path_str)

    print(f"SERVICE (doc_formatter_v1): Attempting dummy formatting for input: '{input_path}'")
    print(f"SERVICE (doc_formatter_v1): Output will be saved to: '{output_path}'")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    file_content = "Default placeholder content if input file cannot be read properly."
    if not input_path.is_file():
        print(f"SERVICE WARNING (doc_formatter_v1): Input file '{input_path}' not found. Using placeholder content.")
        # For a real service, you'd likely raise FileNotFoundError here:
        # raise FileNotFoundError(f"Input file not found at path: {input_path}")
    else:
        try:
            # Using 'rb' to read as bytes first, then decoding, can be more robust
            # for potentially mixed or unknown encodings if you don't expect pure text.
            # However, since python-docx needs strings, we must decode.
            # errors='replace' or 'backslashreplace' might be safer than 'ignore'
            # if you want to see problematic characters rather than silently dropping them.
            with open(input_path, "r", encoding="utf-8", errors="replace") as f:
                file_content = f.read()
            print(f"SERVICE (doc_formatter_v1): Successfully read content from '{input_path}'.")
        except Exception as e_read:
            print(f"SERVICE ERROR (doc_formatter_v1): Could not read input file '{input_path}': {e_read}. Using placeholder content.")
            # For a real service, raise IOError:
            # raise IOError(f"Could not read input file '{input_path}': {e_read}") from e_read
    
    try:
        doc = Document() 
        doc.add_heading('Dummy Formatted Document', level=1)
        
        p1 = doc.add_paragraph()
        run = p1.add_run('This document was "formatted" by DocuMorph AI\'s dummy formatter.')
        run.bold = True
        # It's safer to add text to a run, then style the run,
        # Than to chain .add_run(...).bold = True, though it often works.

        doc.add_paragraph(f'Original file processed: {input_path.name}') # Filename should be safe enough
        
        doc.add_heading('Original Content Snippet:', level=2)
        
        # Sanitize file_content for python-docx to prevent the XML error
        # This is a basic approach. A more robust one might involve a whitelist of allowed characters
        # or more aggressive cleaning depending on expected input.
        sanitized_snippet_list = []
        for char in file_content[:500]: # Process only a snippet
            # Allow ASCII printables, tabs, newlines, carriage returns.
            # ord(char) < 32 checks for most control characters.
            # 127 is DEL. Some characters above 127 might also be problematic depending on font/encoding in DOCX.
            # This is a very basic sanitizer.
            if 32 <= ord(char) <= 126 or char in ('\t', '\n', '\r'):
                sanitized_snippet_list.append(char)
            else:
                sanitized_snippet_list.append('?') # Replace problematic char with a placeholder

        sanitized_snippet = "".join(sanitized_snippet_list)
        
        doc.add_paragraph(sanitized_snippet if sanitized_snippet else "No displayable content from original file.")
        
        doc.save(output_path)
        
        print(f"SERVICE (doc_formatter_v1): Dummy DOCX '{output_path.name}' created successfully at '{output_path}'.")
        return str(output_path)

    except Exception as e_docx: # Catch errors specifically from python-docx or file saving
        error_message = f"Error creating DOCX for '{input_path.name}': {str(e_docx)}"
        print(f"SERVICE ERROR (doc_formatter_v1): {error_message}")
        # This is where the original "All strings must be XML compatible" would be caught
        # Re-raise a more generic exception or a custom one for the route to handle
        raise Exception(error_message) from e_docx